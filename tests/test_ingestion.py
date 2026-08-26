from io import BytesIO

import pytest
from docx import Document

from ai.ingestion import DocumentIngester
from app_helpers import meaningful_lines, validate_upload


def make_docx() -> bytes:
    document = Document()
    document.add_paragraph("Built a synthetic AI product prototype for a student project.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Role"
    table.cell(0, 1).text = "AI Product Intern"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_ingestion_includes_paragraphs_and_tables():
    text = DocumentIngester().extract_bytes(make_docx(), "sample.docx")
    assert "synthetic AI product prototype" in text
    assert "Role | AI Product Intern" in text


def test_upload_validation_rejects_unknown_extension():
    with pytest.raises(ValueError, match="unsupported"):
        validate_upload("resume.exe", b"not an executable")


def test_meaningful_lines_filters_document_instructions():
    text = """How to use this document
Built a working prototype and measured task completion across 20 synthetic cases.
1.
"""
    assert meaningful_lines(text) == [
        "Built a working prototype and measured task completion across 20 synthetic cases."
    ]
