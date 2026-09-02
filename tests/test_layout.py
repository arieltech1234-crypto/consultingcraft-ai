from io import BytesIO

from docx import Document

from demo_content import DEMO_PROFILE, DEMO_RESUME_SCHEMA
from layout.generator import CVGenerator
from layout.page_budget import (
    MIN_LINE_SPACING_PT,
    MIN_SECTION_SPACING_PT,
    MIN_SUB_SECTION_SPACING_PT,
    PageBudgetCalculator,
    fit_to_one_page,
)


PREFERENCES = {
    "font_size": 10,
    "side_margin_inches": 0.55,
    "top_margin_inches": 0.5,
    "bottom_margin_inches": 0.5,
    "line_spacing_pt": 11,
    "section_spacing_pt": 6,
    "sub_section_spacing_pt": 2.5,
}


def test_generated_document_is_a4_and_contains_profile():
    structured = {
        "name": DEMO_PROFILE["name"],
        "contact_line": DEMO_PROFILE["email"],
        "resume_schema": DEMO_RESUME_SCHEMA,
    }
    content = CVGenerator().generate_docx_bytes(structured, PREFERENCES)
    document = Document(BytesIO(content))
    section = document.sections[0]
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert abs(section.page_width.mm - 210) < 0.2
    assert abs(section.page_height.mm - 297) < 0.2
    assert DEMO_PROFILE["name"] in text
    assert DEMO_PROFILE["email"] in text
    assert "AI PRODUCT EXPERIENCE" in text
    assert len(content) > 5_000


def test_generated_document_has_no_stray_empty_paragraphs():
    # docxtpl's plain "{% %}" tags only clear their own text, not the
    # paragraph itself, leaving blank rows behind unless the template uses
    # "{%p %}" (paragraph-scoped) tags -- regression check for the fix.
    structured = {
        "name": DEMO_PROFILE["name"],
        "contact_line": DEMO_PROFILE["email"],
        "resume_schema": DEMO_RESUME_SCHEMA,
    }
    content = CVGenerator().generate_docx_bytes(structured, PREFERENCES)
    document = Document(BytesIO(content))
    empty_paragraphs = [p for p in document.paragraphs if not p.text.strip()]
    assert empty_paragraphs == []


def test_generated_document_preserves_ampersand_in_section_names():
    # docxtpl does not auto-escape XML special characters in substituted
    # text -- a raw "&" is silently dropped rather than rendered.
    structured = {
        "name": DEMO_PROFILE["name"],
        "contact_line": DEMO_PROFILE["email"],
        "resume_schema": DEMO_RESUME_SCHEMA,
    }
    content = CVGenerator().generate_docx_bytes(structured, PREFERENCES)
    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "LEADERSHIP & EXTRACURRICULARS" in text


def test_max_words_for_one_line_respects_a_realistic_width():
    calculator = PageBudgetCalculator(PREFERENCES)
    max_words = calculator.max_words_for_one_line()
    fitting_bullet = " ".join(["word"] * max_words)
    assert calculator.wrapped_line_count(fitting_bullet) == 1

    overflowing_bullet = " ".join(["word"] * (max_words * 3))
    assert calculator.wrapped_line_count(overflowing_bullet) > 1

    # Narrower margins leave less usable width, so fewer words fit one line.
    narrow = PageBudgetCalculator({**PREFERENCES, "side_margin_inches": 0.8})
    assert narrow.max_words_for_one_line() <= max_words


def test_target_character_band_matches_the_real_wrap_boundary():
    calculator = PageBudgetCalculator(PREFERENCES)
    min_chars, max_chars = calculator.target_character_band()

    assert 0 < min_chars < max_chars
    assert max_chars == calculator.characters_per_line()

    # A bullet built up to just under the ceiling must still read as one
    # line, and comfortably over it must wrap -- the band's own definition
    # of "fits". Built of short real words (not one giant token), since a
    # single word longer than the line width doesn't force a wrap the way
    # normal text does.
    words = ["word"] * (max_chars // 5)
    under_ceiling = " ".join(words)
    assert len(under_ceiling) <= max_chars
    assert calculator.wrapped_line_count(under_ceiling) == 1

    over_ceiling = " ".join(words * 3)
    assert len(over_ceiling) > max_chars
    assert calculator.wrapped_line_count(over_ceiling) > 1


def test_demo_content_estimates_to_one_page():
    report = PageBudgetCalculator(PREFERENCES).estimate_fit(DEMO_RESUME_SCHEMA)
    assert report["fits_one_page"] is True
    assert report["estimated_pages"] == 1
    assert 0 < report["utilization"] < 1


def test_long_content_is_flagged_as_multiple_pages():
    long_schema = {
        "sections": [
            {
                "section_name": "Experience",
                "entries": [
                    {
                        "header_left": "Some Company",
                        "header_right": "2020 - 2024",
                        "summary": "",
                        "groups": [
                            {
                                "group_name": "",
                                "group_summary": "",
                                "bullets": [
                                    " ".join(["Delivered measurable product impact"] * 18)
                                    for _ in range(35)
                                ],
                            }
                        ],
                    }
                ],
            }
        ]
    }
    report = PageBudgetCalculator(PREFERENCES).estimate_fit(long_schema)
    assert report["fits_one_page"] is False
    assert report["estimated_pages"] > 1


def _entry(bullets, header="Some Company"):
    return {
        "header_left": header,
        "header_right": "2020 - 2024",
        "summary": "",
        "groups": [{"group_name": "", "group_summary": "", "bullets": bullets}],
    }


def test_fit_to_one_page_tightens_spacing_toward_the_floor():
    # Overflows one page; not narrow enough to be fully rescued by spacing
    # alone (line spacing has a safe natural-height floor so it can't clip
    # text), but tightening must still measurably reduce the overflow.
    schema = {"sections": [{"section_name": "Experience", "entries": [
        _entry([" ".join(["Delivered measurable product impact"] * 8) for _ in range(24)])
    ]}]}
    baseline = PageBudgetCalculator(PREFERENCES).estimate_fit(schema)
    assert baseline["fits_one_page"] is False

    tightened_prefs, estimate = fit_to_one_page(schema, PREFERENCES)
    assert tightened_prefs["section_spacing_pt"] == MIN_SECTION_SPACING_PT
    assert tightened_prefs["sub_section_spacing_pt"] == MIN_SUB_SECTION_SPACING_PT
    assert estimate["utilization"] < baseline["utilization"]


def test_fit_to_one_page_converges_using_section_and_sub_section_spacing():
    # Many short entries across several sections: overflows mostly from
    # per-section/per-entry overhead rather than line wraps, so it needs
    # section/sub-section tightening (not just line spacing) to close.
    schema = {"sections": [
        {"section_name": f"Section {i}", "entries": [_entry(["word"] * 11)]}
        for i in range(4)
    ]}
    baseline = PageBudgetCalculator(PREFERENCES).estimate_fit(schema)
    assert baseline["fits_one_page"] is False

    tightened_prefs, estimate = fit_to_one_page(schema, PREFERENCES)
    assert estimate["fits_one_page"] is True
    # Stops tightening as soon as it fits -- doesn't necessarily walk every
    # knob all the way to its absolute floor.
    assert tightened_prefs["line_spacing_pt"] <= PREFERENCES["line_spacing_pt"]


def test_fit_to_one_page_stops_at_floor_for_content_too_long_to_fit():
    schema = {"sections": [{"section_name": "Experience", "entries": [
        _entry([" ".join(["Delivered measurable product impact"] * 18) for _ in range(35)])
    ]}]}
    tightened_prefs, estimate = fit_to_one_page(schema, PREFERENCES)

    assert tightened_prefs["line_spacing_pt"] == MIN_LINE_SPACING_PT
    assert tightened_prefs["section_spacing_pt"] == MIN_SECTION_SPACING_PT
    assert tightened_prefs["sub_section_spacing_pt"] == MIN_SUB_SECTION_SPACING_PT
    assert estimate["fits_one_page"] is False
