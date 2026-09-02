"""Private, in-memory document ingestion for CVs and job descriptions."""

from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path

from groq import Groq


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".png", ".jpg", ".jpeg"}


class DocumentIngester:
    def __init__(self, api_key: str | None = None, vision_model: str | None = None):
        self.api_key = api_key
        self.vision_model = vision_model

    def extract_text(self, file_path: str) -> str:
        """Extract text from a local file without retaining another copy."""
        path = Path(file_path)
        return self.extract_bytes(path.read_bytes(), path.name)

    def extract_bytes(self, content: bytes, filename: str) -> str:
        """Extract text directly from uploaded bytes without writing to disk."""
        suffix = Path(filename).suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {suffix or 'unknown'}")
        if not content:
            raise ValueError("The uploaded file is empty.")

        if suffix == ".pdf":
            return self._extract_from_pdf(content)
        if suffix == ".docx":
            return self._extract_from_docx(content)
        return self._extract_from_image(content, suffix)

    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        text = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            if extracted.strip():
                text.append(extracted)
        return "\n".join(text)

    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        import docx

        document = docx.Document(BytesIO(content))
        lines = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)

    def _extract_from_image(self, content: bytes, suffix: str) -> str:
        if not self.api_key:
            raise ValueError("An API key is required to read image uploads.")

        client = Groq(api_key=self.api_key)
        vision_model = self.vision_model or self._find_vision_model(client)
        if not vision_model:
            raise ValueError(
                "No vision-capable model is available. Upload a PDF or DOCX instead."
            )

        # Downscale the image to drastically reduce base64 token size (TPM limits)
        from PIL import Image
        import io
        
        try:
            img = Image.open(BytesIO(content))
            img.thumbnail((800, 800), Image.Resampling.LANCZOS)
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            buffer = io.BytesIO()
            img.save(buffer, format="JPEG", quality=80)
            content = buffer.getvalue()
            mime_type = "image/jpeg"
        except Exception as e:
            print(f"Failed to downscale image: {e}")
            mime_type = "image/png" if suffix == ".png" else "image/jpeg"

        encoded = base64.b64encode(content).decode("ascii")
        import time
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = client.chat.completions.create(
                    model=vision_model,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": (
                                        "Extract the text from this resume or job description. "
                                        "Return plain text only and preserve section order."
                                    ),
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {"url": f"data:{mime_type};base64,{encoded}"},
                                },
                            ],
                        }
                    ],
                    max_tokens=2048,
                )
                return response.choices[0].message.content or ""
            except Exception as e:
                if "rate_limit_exceeded" in str(e).lower() and "tokens per minute" in str(e).lower():
                    # No amount of retrying fixes a payload that is structurally too large for TPM
                    raise ValueError(f"Image is still too large for Groq's TPM limits. Please upload a PDF instead. Error: {e}")
                if attempt < max_retries - 1:
                    time.sleep(5 * (attempt + 1))
                else:
                    raise ValueError(f"Failed to process image after {max_retries} attempts: {e}")
        return ""

    # Known-good vision models, most preferred first. Falls back to a
    # keyword match only when none of these are available for the key.
    PREFERRED_VISION_MODELS = [
        "meta-llama/llama-4-scout-17b-16e-instruct",
        "meta-llama/llama-4-maverick-17b-128e-instruct",
    ]

    @classmethod
    def _find_vision_model(cls, client: Groq) -> str | None:
        models = client.models.list().data
        available = {model.id for model in models}

        for preferred in cls.PREFERRED_VISION_MODELS:
            if preferred in available:
                return preferred

        candidates = [
            model_id
            for model_id in available
            if any(token in model_id.lower() for token in ("vision", "vl", "scout", "qwen"))
        ]
        return sorted(candidates)[0] if candidates else None
