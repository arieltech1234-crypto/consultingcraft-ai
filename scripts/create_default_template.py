"""
Generates a default consulting CV template as a .docx file.
"""
import os
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_default_template(output_path: str):
    doc = docx.Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def add_section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    def add_placeholder_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run("YOUR NAME")
    name_run.bold = True
    name_run.font.name = 'Arial'
    name_run.font.size = Pt(14)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = contact.add_run("email@example.com | +1-234-567-8900 | City, Country")
    c_run.font.name = 'Arial'
    c_run.font.size = Pt(10)

    # Sections
    add_section_header("EDUCATION")
    add_placeholder_bullet("[Degree, University, Year — GPA/Honors]")
    add_placeholder_bullet("[Relevant coursework or academic achievements]")

    add_section_header("WORK EXPERIENCE")
    add_placeholder_bullet("[Result-Action-Context bullet for work experience]")
    add_placeholder_bullet("[Result-Action-Context bullet for work experience]")

    add_section_header("EXTRACURRICULARS & LEADERSHIP")
    add_placeholder_bullet("[Leadership role, organization, impact]")
    add_placeholder_bullet("[Volunteer work, club involvement, achievements]")

    add_section_header("SKILLS & INTERESTS")
    add_placeholder_bullet("[Technical skills, languages, certifications]")
    add_placeholder_bullet("[Personal interests, hobbies]")

    doc.save(output_path)
    print(f"Default template created at: {output_path}")

if __name__ == "__main__":
    os.makedirs("data/templates", exist_ok=True)
    create_default_template("data/templates/Default_Consulting_Template.docx")
