"""Builds the Jinja-tagged DOCX template used for every export.

The visual chrome here (colors, fonts, margins, tab positions) is copied
directly from the user's real, working reference resume
(Anuraag_Consulting_CV_Final_v2.docx) -- inspected via its raw OOXML -- so
the export actually matches that document instead of an approximation.
"""
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_LINE_SPACING


def set_solid_shading(paragraph, hex_fill: str) -> None:
    """Apply a solid background color behind a paragraph. python-docx has no
    high-level API for this -- it's a direct OOXML <w:shd> element."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


SECTION_FILL = "153D63"  # dark navy, matches the reference's section headers
ENTRY_FILL = "D1D1D1"    # light gray, matches the reference's entry headers
FONT_SIZE_PT = 9         # matches the reference (not 10)

doc = docx.Document()
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(FONT_SIZE_PT)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = Pt(1)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST

bullet_style = doc.styles["List Bullet"]
bullet_style.font.name = "Arial"
bullet_style.font.size = Pt(FONT_SIZE_PT)
bullet_style.paragraph_format.space_before = Pt(0)
bullet_style.paragraph_format.space_after = Pt(0)
bullet_style.paragraph_format.line_spacing = Pt(1)
bullet_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST

# Document margins and page size (A4). Matches the reference's tight
# top/bottom margins (~0.2in) -- layout_preferences overrides all of this
# per-export via CVGenerator._apply_layout.
for section in doc.sections:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

TAB_RIGHT = Inches(7.4)

# Header: name + contact line
head_p = doc.add_paragraph()
head_p.paragraph_format.tab_stops.add_tab_stop(TAB_RIGHT, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
head_p.paragraph_format.space_after = Pt(4)
head_r = head_p.add_run("{{ name }}")
head_r.bold = True
head_r.font.size = Pt(FONT_SIZE_PT)
contact_r = head_p.add_run("\t{{ contact_line }}")
contact_r.bold = False
contact_r.font.size = Pt(FONT_SIZE_PT)

# "{%p ... %}" (rather than plain "{% ... %}") tells docxtpl to remove the
# ENTIRE paragraph the tag lives in once rendered. Plain "{% %}" tags only
# clear their own text and were leaving the paragraph itself behind empty --
# confirmed directly: a rendered document had 3 blank paragraphs stacked in
# a row where a "{% if %}/{% endif %}" pair used to be.
doc.add_paragraph("{%p for section in sections %}")

# Section name -- solid dark-navy bar with white text, matching "EDUCATION" etc.
sec_p = doc.add_paragraph()
sec_p.paragraph_format.tab_stops.add_tab_stop(TAB_RIGHT, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
set_solid_shading(sec_p, SECTION_FILL)
sec_r = sec_p.add_run("{{ section.section_name }}")  # already uppercased in CVGenerator, before XML-escaping
sec_r.bold = True
sec_r.font.size = Pt(FONT_SIZE_PT)
sec_r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph("{%p for entry in section.entries %}")

# Entry header -- solid light-gray bar, left header bold + right-aligned date
ent_p = doc.add_paragraph()
ent_p.paragraph_format.tab_stops.add_tab_stop(TAB_RIGHT, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
set_solid_shading(ent_p, ENTRY_FILL)
el = ent_p.add_run("{{ entry.header_left }}")
el.bold = True
el.font.size = Pt(FONT_SIZE_PT)
er = ent_p.add_run("\t{{ entry.header_right }}")
er.bold = False
er.font.size = Pt(FONT_SIZE_PT)

doc.add_paragraph("{%p if entry.summary %}")
sum_p = doc.add_paragraph()
sum_r = sum_p.add_run("{{ entry.summary }}")
sum_r.italic = True
sum_r.font.size = Pt(FONT_SIZE_PT)
doc.add_paragraph("{%p endif %}")

doc.add_paragraph("{%p for group in entry.groups %}")
doc.add_paragraph("{%p if group.group_name %}")
grp_p = doc.add_paragraph()
grp_r = grp_p.add_run("{{ group.group_name }}")
grp_r.underline = True
grp_r.font.size = Pt(FONT_SIZE_PT)
doc.add_paragraph("{%p endif %}")

doc.add_paragraph("{%p for bullet in group.bullets %}")
bul_p = doc.add_paragraph(style="List Bullet")
bul_r = bul_p.add_run("{{ bullet }}")
bul_r.font.size = Pt(FONT_SIZE_PT)
doc.add_paragraph("{%p endfor %}")
doc.add_paragraph("{%p endfor %}")
doc.add_paragraph("{%p endfor %}")
doc.add_paragraph("{%p endfor %}")

doc.save("data/templates/isb_jinja_template.docx")
print("Template created.")
